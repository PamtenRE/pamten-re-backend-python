# backend/document_generator.py
import io, base64, re, os
from bs4 import BeautifulSoup
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

from docx import Document
from docx.shared import Pt, Inches, RGBColor

# ------------------------------------------------------------
# Helpers (sanitization, minimal HTML -> DOCX runs, safe images)
# ------------------------------------------------------------

# Remove control chars that can crash python-docx
CONTROL = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')

def _clean(s):
    if s is None:
        return ""
    return CONTROL.sub("", str(s)).strip()

def _add_b64_image(paragraph, b64, width_inches=1.3, name="image.png"):
    """Add a base64 image to a paragraph; skip silently if corrupted."""
    if not b64:
        return
    try:
        if b64.startswith("data:"):
            b64 = b64.split(",", 1)[1]
        bio = io.BytesIO(base64.b64decode(b64, validate=True))
        bio.name = name  # lets python-docx infer image type
        paragraph.add_run().add_picture(bio, width=Inches(width_inches))
    except Exception:
        # Skip bad/corrupt images instead of crashing generation
        pass

def _inline_html(p, html):
    """Very small subset of inline HTML -> DOCX runs."""
    if not html:
        return
    soup = BeautifulSoup(html, "html.parser")
    for n in soup.contents:
        tag = getattr(n, "name", None)
        if tag == "strong":
            r = p.add_run(n.get_text()); r.bold = True
        elif tag == "em":
            r = p.add_run(n.get_text()); r.italic = True
        elif tag == "br":
            p.add_run().add_break()
        elif tag == "a":
            r = p.add_run(n.get_text()); r.font.underline = True; r.font.color.rgb = RGBColor(0,0,255)
        elif isinstance(n, str):
            p.add_run(n)
        elif getattr(n, "contents", None):
            _inline_html(p, str(n))

# ------------------------------------------------------------
# DOCX GENERATION (returns BytesIO buffer, not a Document)
# ------------------------------------------------------------
def generate_docx_from_data(data: dict) -> io.BytesIO:
    data = data or {}
    doc = Document()

    # Styles
    style = data.get("styleOptions", {}) or {}
    font_name = (style.get("fontFamily", "Calibri").split(",")[0]).strip()
    font_size = style.get("fontSize", 11)
    # RGBColor.from_string expects hex without '#'
    accent = RGBColor.from_string(style.get("accentColor", "#34495e").lstrip("#"))

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)

    # Header table (logo | name+contact | profile)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    left, mid, right = table.rows[0].cells

    if data.get("pamtenLogoBase64"):
        _add_b64_image(left.paragraphs[0], data["pamtenLogoBase64"], 1.2, "logo.png")

    person = data.get("personal", {}) or {}
    p = mid.paragraphs[0]
    r = p.add_run(_clean(person.get("name"))); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = accent

    contacts = [
        _clean(person.get("email")),
        _clean(person.get("phone")),
        _clean(person.get("location")),
    ]
    legal = _clean(person.get("legalStatus"))
    if legal and legal.lower() != "prefer not to say":
        contacts.append(legal)
    mid.add_paragraph(" | ".join([c for c in contacts if c]))

    if data.get("profilePicBase64"):
        _add_b64_image(right.paragraphs[0], data["profilePicBase64"], 1.1, "profile.jpg")

    doc.add_paragraph()  # spacer

    # Summary
    if data.get("summary"):
        doc.add_heading("Summary", level=2)
        p = doc.add_paragraph()
        _inline_html(p, data.get("summary"))

    # Generic sections
    for section, hdr in [
        ("experience", "Experience"),
        ("education", "Education"),
        ("skills", "Skills"),
        ("projects", "Projects"),
        ("publications", "Publications"),
        ("certifications", "Certifications"),
    ]:
        items = data.get(section) or []
        # If it's not a list of dicts (except skills), skip
        if not any(isinstance(x, dict) for x in items) and section != "skills":
            continue

        doc.add_heading(hdr, level=2)

        if section == "skills":
            for s in items:
                p = doc.add_paragraph()
                p.add_run(_clean(s.get("category")) + ": ").bold = True
                p.add_run(_clean(s.get("skills_list")))
            continue

        for it in items:
            if not isinstance(it, dict):
                continue
            p = doc.add_paragraph()
            if section == "experience":
                p.add_run(_clean(it.get("jobTitle"))).bold = True
                p.add_run(f"\n{_clean(it.get('company'))} | {_clean(it.get('dates'))}\n").italic = True
                _inline_html(p, it.get("description"))
            elif section == "education":
                p.add_run(_clean(it.get("degree"))).bold = True
                p.add_run(f", {_clean(it.get('institution'))}\n")
                tail = _clean(it.get("graduationYear"))
                if it.get("gpa"): tail += f" | GPA: {_clean(it.get('gpa'))}"
                p.add_run(tail).italic = True
            elif section == "projects":
                p.add_run(_clean(it.get("title"))).bold = True
                p.add_run(f" ({_clean(it.get('date'))})\n").italic = True
                _inline_html(p, it.get("description"))
            elif section == "publications":
                p.add_run(_clean(it.get("title"))).bold = True
                p.add_run(f" ({_clean(it.get('date'))})\n").italic = True
                p.add_run(f"{_clean(it.get('authors'))} - {_clean(it.get('journal'))}")
            elif section == "certifications":
                p.add_run(_clean(it.get("name"))).bold = True
                issuer = _clean(it.get("issuer"))
                if it.get("date"): issuer += f" | {_clean(it.get('date'))}"
                p.add_run(f"\n{issuer}").italic = True

    # Serialize to BytesIO for Flask send_file
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    out.name = "resume.docx"
    return out

# ------------------------------------------------------------
# PDF GENERATION (unchanged pipeline you already use)
# ------------------------------------------------------------
def clean_text(text: str) -> str:
    if not text:
        return ""
    cleaned_text = re.sub(r"[ \t]+", " ", text)
    cleaned_text = re.sub(r"\n{2,}", "\n", cleaned_text)
    cleaned_text = "\n".join(line.strip() for line in cleaned_text.split("\n"))
    return cleaned_text.strip()

def generate_pdf_from_data(data):
    data = data or {}
    data["pamtenLogoSrc"] = data.get("pamtenLogoBase64") or None
    data["profilePicSrc"] = data.get("profilePicBase64") or None

    if data.get("skills"):
        for skill in data["skills"]:
            skill["skills_list"] = clean_text(skill.get("skills_list", ""))

    env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(__file__), "assets")))
    template = env.get_template("resume_template.html")
    rendered_html = template.render(**data)
    return HTML(string=rendered_html).write_pdf()
